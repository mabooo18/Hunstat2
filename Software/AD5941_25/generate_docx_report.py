# Python script to convert TECHNICAL_REPORT_EN.md into a beautifully formatted Word Document (.docx).

import re
import os
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

def set_cell_background(cell, color_hex):
    """Sets background color of a table cell."""
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    """Sets internal padding (margins) of a cell in dxa (1/20 pt)."""
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('w:top', top), ('w:bottom', bottom), ('w:left', left), ('w:right', right)]:
        node = OxmlElement(m)
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def parse_formatting(paragraph, text):
    """Parses inline bold (**text**) and converts to word runs."""
    pattern = re.compile(r'(\*\*.*?\*\*|\$.*?\$|`.*?`)')
    parts = pattern.split(text)
    
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith('$') and part.endswith('$'):
            # Simple math notation wrapper
            run = paragraph.add_run(part[1:-1])
            run.italic = True
        elif part.startswith('`') and part.endswith('`'):
            run = paragraph.add_run(part[1:-1])
            run.font.name = 'Courier New'
            run.font.size = Pt(9.5)
            run.font.color.rgb = RGBColor(0x8F, 0x1F, 0x3F)
        else:
            paragraph.add_run(part)

def main():
    md_path = "TECHNICAL_REPORT_EN.md"
    docx_path = "TECHNICAL_REPORT_EN.docx"
    
    if not os.path.exists(md_path):
        print(f"Error: {md_path} tidak ditemukan.")
        return 1

    doc = docx.Document()
    
    # 1. Page Margins Setup
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)
        
    # 2. Configure Typography Styles
    # Normal Text (Body)
    style_normal = doc.styles['Normal']
    font_normal = style_normal.font
    font_normal.name = 'Times New Roman'
    font_normal.size = Pt(11)
    font_normal.color.rgb = RGBColor(0x22, 0x22, 0x22)
    
    # Read Markdown Lines
    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()
        
    in_code_block = False
    code_text = []
    
    in_table = False
    table_rows = []
    
    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()
        
        # --- Handle Code Blocks ---
        if stripped.startswith("```"):
            if in_code_block:
                # End of code block, write to document
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Inches(0.4)
                p.paragraph_format.right_indent = Inches(0.4)
                run = p.add_run("\n".join(code_text))
                run.font.name = 'Courier New'
                run.font.size = Pt(9.0)
                run.font.color.rgb = RGBColor(0x33, 0x41, 0x55)
                # Set light grey background
                pBdr = parse_xml(r'<w:pBdr xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
                                 r'<w:left w:val="single" w:sz="24" w:space="8" w:color="38BDF8"/>'
                                 r'</w:pBdr>')
                p._p.get_or_add_pPr().append(pBdr)
                shd = parse_xml(r'<w:shd xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" w:fill="F8FAFC"/>')
                p._p.get_or_add_pPr().append(shd)
                
                in_code_block = False
                code_text = []
            else:
                in_code_block = True
            i += 1
            continue
            
        if in_code_block:
            code_text.append(line.rstrip('\r\n'))
            i += 1
            continue
            
        # --- Handle Tables ---
        if stripped.startswith("|"):
            in_table = True
            table_rows.append(stripped)
            i += 1
            continue
        elif in_table:
            # End of table block, render table
            # Parse table rows
            valid_rows = []
            for r in table_rows:
                parts = [p.strip() for p in r.split("|")[1:-1]]
                if all(p == '' or all(c == '-' for c in p) for p in parts):
                    # Separator row like |---|---|
                    continue
                valid_rows.append(parts)
                
            if valid_rows:
                num_cols = len(valid_rows[0])
                table = doc.add_table(rows=len(valid_rows), cols=num_cols)
                table.style = 'Table Grid'
                table.autofit = True
                
                for r_idx, row_data in enumerate(valid_rows):
                    for c_idx, val in enumerate(row_data):
                        cell = table.cell(r_idx, c_idx)
                        cell.text = "" # Clear default text
                        p = cell.paragraphs[0]
                        p.paragraph_format.space_before = Pt(4)
                        p.paragraph_format.space_after = Pt(4)
                        p.paragraph_format.line_spacing = 1.05
                        
                        run = p.add_run(val)
                        run.font.name = 'Arial'
                        run.font.size = Pt(9.5)
                        
                        set_cell_margins(cell, top=100, bottom=100, left=150, right=150)
                        
                        if r_idx == 0:
                            # Header styling
                            run.bold = True
                            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                            set_cell_background(cell, "1E293B") # Dark slate
                        else:
                            # Alternating background colors
                            if r_idx % 2 == 0:
                                set_cell_background(cell, "F1F5F9") # Soft grey
                            else:
                                set_cell_background(cell, "FFFFFF")
                                
                doc.add_paragraph() # Add spacer after table
            in_table = False
            table_rows = []
            
        # --- Handle Headings ---
        if stripped.startswith("# "):
            title_text = stripped[2:]
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(24)
            p.paragraph_format.space_after = Pt(12)
            run = p.add_run(title_text)
            run.font.name = 'Arial'
            run.font.size = Pt(22)
            run.font.bold = True
            run.font.color.rgb = RGBColor(0x1E, 0x29, 0x3B)
            i += 1
            continue
        elif stripped.startswith("## "):
            h_text = stripped[3:]
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(18)
            p.paragraph_format.space_after = Pt(6)
            p.paragraph_format.keep_with_next = True
            run = p.add_run(h_text)
            run.font.name = 'Arial'
            run.font.size = Pt(14)
            run.font.bold = True
            run.font.color.rgb = RGBColor(0x0F, 0x76, 0x6E) # Teal color
            i += 1
            continue
        elif stripped.startswith("### "):
            h_text = stripped[4:]
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(4)
            p.paragraph_format.keep_with_next = True
            run = p.add_run(h_text)
            run.font.name = 'Arial'
            run.font.size = Pt(11.5)
            run.font.bold = True
            run.font.color.rgb = RGBColor(0x1E, 0x29, 0x3B)
            i += 1
            continue
            
        # --- Handle Lists (Bullet Points) ---
        if stripped.startswith("* ") or stripped.startswith("- "):
            bullet_text = stripped[2:]
            p = doc.add_paragraph(style='List Bullet')
            p.paragraph_format.space_after = Pt(3)
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.line_spacing = 1.15
            parse_formatting(p, bullet_text)
            i += 1
            continue
            
        # --- Handle Numbered Lists ---
        if re.match(r'^\d+\.\s', stripped):
            num_text = re.sub(r'^\d+\.\s', '', stripped)
            p = doc.add_paragraph(style='List Number')
            p.paragraph_format.space_after = Pt(4)
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.line_spacing = 1.15
            parse_formatting(p, num_text)
            i += 1
            continue
            
        # --- Handle Blank Lines ---
        if not stripped:
            i += 1
            continue
            
        # --- Handle Separator Lines ---
        if stripped == "---":
            # Add a thin grey bottom border to simulate horizontal line
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(12)
            pBdr = parse_xml(r'<w:pBdr xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
                             r'<w:bottom w:val="single" w:sz="6" w:space="1" w:color="CBD5E1"/>'
                             r'</w:pBdr>')
            p._p.get_or_add_pPr().append(pBdr)
            i += 1
            continue
            
        # --- Handle Blockquotes ---
        if stripped.startswith(">"):
            bq_text = re.sub(r'^>\s*', '', stripped)
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.4)
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(6)
            
            # Left border and light background
            pBdr = parse_xml(r'<w:pBdr xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
                             r'<w:left w:val="single" w:sz="18" w:space="8" w:color="94A3B8"/>'
                             r'</w:pBdr>')
            p._p.get_or_add_pPr().append(pBdr)
            
            run = p.add_run(bq_text)
            run.font.italic = True
            run.font.color.rgb = RGBColor(0x47, 0x55, 0x69)
            i += 1
            continue

        # --- Handle Normal Paragraphs ---
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.line_spacing = 1.15
        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        parse_formatting(p, stripped)
        i += 1
        
    doc.save(docx_path)
    print(f"Sukses mengonversi dokumen Word ke {docx_path}")
    return 0

if __name__ == "__main__":
    import sys
    sys.exit(main())
